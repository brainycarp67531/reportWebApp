from flask import Flask, render_template_string, request, send_file
from docx import Document
import fitz  # PyMuPDF

app = Flask(__name__)

HTML_TEMPLATE = '''
<!doctype html>
<title>Material and Weight PDF Generator</title>
<h2>Enter Material and Weight</h2>
<form method=post>
  Material: <input type=text name=material><br><br>
  Weight: <input type=text name=weight><br><br>
  <input type=submit value=Generate>
</form>
{% if pdf_ready %}
  <p>PDF generated successfully. /downloadDownload PDF</a></p>
{% endif %}
'''

@app.route('/', methods=['GET', 'POST'])
def index():
    pdf_ready = False
    if request.method == 'POST':
        material = request.form['material']
        weight = request.form['weight']

        if not material or not weight:
            return render_template_string(HTML_TEMPLATE, pdf_ready=False)

        # Load and modify the Word template
        doc = Document("template.docx")
        for para in doc.paragraphs:
            if "{{material}}" in para.text:
                para.text = para.text.replace("{{material}}", material)
            if "{{weight}}" in para.text:
                para.text = para.text.replace("{{weight}}", weight)

        doc.save("output.docx")

        # Create PDF using PyMuPDF
        pdf_doc = fitz.open()
        for para in doc.paragraphs:
            page = pdf_doc.new_page()
            page.insert_text((72, 72), para.text, fontsize=12)

        pdf_doc.save("output.pdf")
        pdf_doc.close()

        pdf_ready = True

    return render_template_string(HTML_TEMPLATE, pdf_ready=pdf_ready)

@app.route('/download')
def download():
    return send_file("output.pdf", as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
